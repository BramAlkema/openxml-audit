"""Property-element ordering scenarios for the Word roundtrip oracle.

Builds DOCX files using python-docx (so the document has a real Word-
template foundation: full styles, fonts, relationships) and then mutates
the property-element children via lxml to test specific orderings.

Each `OrderingMatrix` describes one Word property complex type: the
canonical child order, the minimum-valid attributes to attach to each
child so Word opens the file cleanly, and a host materializer that
embeds the property element in a DOCX shape Word will accept.

Spec: `specs/011-word-roundtrip-oracle.md` Phase 2.
"""

from __future__ import annotations

from collections.abc import Callable
from dataclasses import dataclass, field
from pathlib import Path

from docx import Document
from lxml import etree
from tools.oracle.diff import slugify_scenario_id

from openxml_audit.namespaces import WORDPROCESSINGML

W = WORDPROCESSINGML
WNS = f"{{{W}}}"


@dataclass(frozen=True)
class ScenarioSpec:
    """One ordering scenario to materialize and roundtrip."""

    id: str
    parent_local: str  # e.g. "trPr"
    input_children: tuple[str, ...]  # local names in the order to test
    description: str


@dataclass(frozen=True)
class OrderingMatrix:
    """Configuration bundle for one property complex type.

    Used by the spec-010 driver to generate scenarios, materialize DOCX
    files for them, and label the resulting oracle baseline.
    """

    parent_local: str  # e.g. "trPr", "tblPr", "tcPr", "sectPr"
    spec_section: str  # ECMA-376 reference, e.g. "ECMA-376 §17.4.79"
    canonical_children: tuple[str, ...]
    child_attrs: dict[str, dict[str, str]] = field(default_factory=dict)
    materialize: Callable[[ScenarioSpec, OrderingMatrix, Path], Path] = field(
        default=lambda spec, matrix, dest: dest
    )


# --- generic scenario builders --------------------------------------------


def baseline_scenario(matrix: OrderingMatrix) -> ScenarioSpec:
    """Canonical-order property — the control case. Preserved → engine works."""
    return ScenarioSpec(
        id=slugify_scenario_id(matrix.parent_local, "baseline"),
        parent_local=matrix.parent_local,
        input_children=tuple(matrix.canonical_children),
        description=f"Canonical {matrix.parent_local} child order (control)",
    )


def pairwise_swap_scenarios(matrix: OrderingMatrix) -> list[ScenarioSpec]:
    """Every pairwise swap among canonical children. For n canonical
    children this produces n*(n-1)/2 scenarios."""
    canonical = matrix.canonical_children
    scenarios: list[ScenarioSpec] = []
    for i in range(len(canonical)):
        for j in range(i + 1, len(canonical)):
            ordered = list(canonical)
            ordered[i], ordered[j] = ordered[j], ordered[i]
            scenarios.append(
                ScenarioSpec(
                    id=slugify_scenario_id(
                        matrix.parent_local, "swap", canonical[i], canonical[j]
                    ),
                    parent_local=matrix.parent_local,
                    input_children=tuple(ordered),
                    description=(
                        f"Swap '{canonical[i]}' and '{canonical[j]}' in {matrix.parent_local}"
                    ),
                )
            )
    return scenarios


def full_reverse_scenario(matrix: OrderingMatrix) -> ScenarioSpec:
    """Canonical order fully reversed. Maximum disturbance — should
    trigger any ordering enforcement Word has."""
    return ScenarioSpec(
        id=slugify_scenario_id(matrix.parent_local, "reverse"),
        parent_local=matrix.parent_local,
        input_children=tuple(reversed(matrix.canonical_children)),
        description=f"{matrix.parent_local} children fully reversed",
    )


def all_scenarios(matrix: OrderingMatrix) -> list[ScenarioSpec]:
    """Baseline + all pairwise swaps + full reverse — the standard matrix."""
    return [
        baseline_scenario(matrix),
        *pairwise_swap_scenarios(matrix),
        full_reverse_scenario(matrix),
    ]


def _set_child_attrs(
    child: etree._Element,
    child_local: str,
    attrs: dict[str, dict[str, str]],
) -> None:
    """Apply the minimum-valid attribute set to a freshly-created child."""
    for attr_local, attr_val in attrs.get(child_local, {}).items():
        child.set(f"{WNS}{attr_local}", attr_val)


# --- CT_TrPr ---------------------------------------------------------------

TRPR_CHILD_ATTRS: dict[str, dict[str, str]] = {
    "divId": {"val": "0"},
    "gridBefore": {"val": "0"},
    "gridAfter": {"val": "0"},
    "cnfStyle": {"val": "000000000000"},
    "wBefore": {"w": "0", "type": "dxa"},
    "wAfter": {"w": "0", "type": "dxa"},
    "tblCellSpacing": {"w": "0", "type": "dxa"},
    "jc": {"val": "center"},
    "trHeight": {},
    "hidden": {},
    "cantSplit": {},
    "tblHeader": {},
}


def materialize_trpr_scenario(
    spec: ScenarioSpec, matrix: OrderingMatrix, dest: Path
) -> Path:
    """Embed a synthetic w:trPr inside a one-row table on a python-docx
    Document. Children get attributes from `matrix.child_attrs` so Word
    opens the file cleanly."""
    doc = Document()
    table = doc.add_table(rows=1, cols=2)
    tr = table.rows[0]._tr  # noqa: SLF001 — python-docx exposes the lxml element this way

    for existing in tr.findall(f"{WNS}trPr"):
        tr.remove(existing)

    trPr = etree.Element(f"{WNS}trPr")
    tr.insert(0, trPr)
    for child_local in spec.input_children:
        child = etree.SubElement(trPr, f"{WNS}{child_local}")
        _set_child_attrs(child, child_local, matrix.child_attrs)

    dest.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(dest))
    return dest


TRPR_MATRIX = OrderingMatrix(
    parent_local="trPr",
    spec_section="ECMA-376 §17.4.79",
    canonical_children=(
        "cnfStyle", "divId", "gridBefore", "gridAfter", "wBefore", "wAfter",
        "trHeight", "hidden", "cantSplit", "tblHeader", "tblCellSpacing", "jc",
    ),
    child_attrs=TRPR_CHILD_ATTRS,
    materialize=materialize_trpr_scenario,
)

# Backwards-compat name kept for the existing test suite.
TRPR_CANONICAL_CORE = TRPR_MATRIX.canonical_children


def trpr_baseline(canonical: tuple[str, ...]) -> ScenarioSpec:
    return ScenarioSpec(
        id=slugify_scenario_id("trPr", "baseline"),
        parent_local="trPr",
        input_children=tuple(canonical),
        description="Canonical CT_TrPr child order (control)",
    )


def trpr_pairwise_swaps(canonical: tuple[str, ...]) -> list[ScenarioSpec]:
    return pairwise_swap_scenarios(
        OrderingMatrix(
            parent_local="trPr",
            spec_section=TRPR_MATRIX.spec_section,
            canonical_children=canonical,
        )
    )


def trpr_full_reverse(canonical: tuple[str, ...]) -> ScenarioSpec:
    return ScenarioSpec(
        id=slugify_scenario_id("trPr", "reverse"),
        parent_local="trPr",
        input_children=tuple(reversed(canonical)),
        description="CT_TrPr children fully reversed",
    )


# --- CT_TblPr --------------------------------------------------------------

TBLPR_CHILD_ATTRS: dict[str, dict[str, str]] = {
    "tblStyle": {"val": "Normal"},          # references a styleId; Normal is universally present
    "tblOverlap": {"val": "never"},          # ST_TblOverlap: never|overlap
    "bidiVisual": {},                         # OnOff optional
    "tblW": {"w": "0", "type": "auto"},
    "jc": {"val": "center"},                 # ST_JcTable
    "tblCellSpacing": {"w": "0", "type": "dxa"},
    "tblInd": {"w": "0", "type": "dxa"},
    "tblBorders": {},                         # container with optional children
    "shd": {"val": "clear", "color": "auto", "fill": "auto"},
    "tblLayout": {"type": "autofit"},        # ST_TblLayoutType
    "tblCellMar": {},                         # container
    "tblLook": {"val": "0000"},               # 16-bit hex bitstring
    "tblCaption": {"val": "test"},
    "tblDescription": {"val": "test"},
}


def materialize_tblpr_scenario(
    spec: ScenarioSpec, matrix: OrderingMatrix, dest: Path
) -> Path:
    """Embed a synthetic w:tblPr at the start of a table body on a
    python-docx Document. Children get attributes from `matrix.child_attrs`."""
    doc = Document()
    table = doc.add_table(rows=1, cols=2)
    tbl = table._tbl  # noqa: SLF001

    for existing in tbl.findall(f"{WNS}tblPr"):
        tbl.remove(existing)

    tblPr = etree.Element(f"{WNS}tblPr")
    tbl.insert(0, tblPr)
    for child_local in spec.input_children:
        child = etree.SubElement(tblPr, f"{WNS}{child_local}")
        _set_child_attrs(child, child_local, matrix.child_attrs)

    dest.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(dest))
    return dest


# CT_TblPr canonical core. Skips tblpPr (complex positioning attrs that
# can interact with the rest of the table) and tblPrChange (track-change).
TBLPR_MATRIX = OrderingMatrix(
    parent_local="tblPr",
    spec_section="ECMA-376 §17.4.60",
    canonical_children=(
        "tblStyle", "tblOverlap", "bidiVisual", "tblW", "jc", "tblCellSpacing",
        "tblInd", "tblBorders", "shd", "tblLayout", "tblCellMar", "tblLook",
        "tblCaption", "tblDescription",
    ),
    child_attrs=TBLPR_CHILD_ATTRS,
    materialize=materialize_tblpr_scenario,
)


# --- CT_TcPr ---------------------------------------------------------------

TCPR_CHILD_ATTRS: dict[str, dict[str, str]] = {
    "cnfStyle": {"val": "000000000000"},
    "tcW": {"w": "0", "type": "auto"},
    "gridSpan": {"val": "1"},
    "hMerge": {"val": "continue"},           # ST_Merge
    "vMerge": {"val": "continue"},
    "tcBorders": {},
    "shd": {"val": "clear", "color": "auto", "fill": "auto"},
    "noWrap": {},
    "tcMar": {},
    "textDirection": {"val": "lrTb"},        # ST_TextDirection
    "tcFitText": {},
    "vAlign": {"val": "top"},                # ST_VerticalJc
    "hideMark": {},
}


def materialize_tcpr_scenario(
    spec: ScenarioSpec, matrix: OrderingMatrix, dest: Path
) -> Path:
    """Embed a synthetic w:tcPr at the start of a table cell on a
    python-docx Document."""
    doc = Document()
    table = doc.add_table(rows=1, cols=2)
    tc = table.rows[0].cells[0]._tc  # noqa: SLF001

    for existing in tc.findall(f"{WNS}tcPr"):
        tc.remove(existing)

    tcPr = etree.Element(f"{WNS}tcPr")
    tc.insert(0, tcPr)
    for child_local in spec.input_children:
        child = etree.SubElement(tcPr, f"{WNS}{child_local}")
        _set_child_attrs(child, child_local, matrix.child_attrs)

    dest.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(dest))
    return dest


# CT_TcPr canonical core. Skips cellIns/cellDel/cellMerge/tcPrChange
# (track-change and complex merge-state children).
TCPR_MATRIX = OrderingMatrix(
    parent_local="tcPr",
    spec_section="ECMA-376 §17.4.70",
    canonical_children=(
        "cnfStyle", "tcW", "gridSpan", "hMerge", "vMerge", "tcBorders",
        "shd", "noWrap", "tcMar", "textDirection", "tcFitText", "vAlign",
        "hideMark",
    ),
    child_attrs=TCPR_CHILD_ATTRS,
    materialize=materialize_tcpr_scenario,
)


# --- CT_SectPr -------------------------------------------------------------

SECTPR_CHILD_ATTRS: dict[str, dict[str, str]] = {
    "footnotePr": {},
    "endnotePr": {},
    "type": {"val": "nextPage"},             # ST_SectionMark
    "pgSz": {"w": "12240", "h": "15840"},    # US Letter
    "pgMar": {
        "top": "1440", "right": "1440", "bottom": "1440", "left": "1440",
        "header": "720", "footer": "720", "gutter": "0",
    },
    "pgBorders": {},
    "lnNumType": {"countBy": "1"},           # ST_DecimalNumber-typed
    "pgNumType": {"fmt": "decimal"},         # ST_NumberFormat
    "cols": {"num": "1", "space": "720"},
    "formProt": {"val": "false"},
    "vAlign": {"val": "top"},
    "noEndnote": {},
    "titlePg": {},
    "textDirection": {"val": "lrTb"},
    "bidi": {},
    "rtlGutter": {},
    "docGrid": {"type": "default", "linePitch": "360"},
    "paperSrc": {"first": "0", "other": "0"},
}


def materialize_sectpr_scenario(
    spec: ScenarioSpec, matrix: OrderingMatrix, dest: Path
) -> Path:
    """Replace the body's sectPr (python-docx adds a default one) with a
    synthetic one whose children appear in the requested order."""
    doc = Document()
    body = doc.element.body  # noqa: SLF001

    for existing in body.findall(f"{WNS}sectPr"):
        body.remove(existing)

    sectPr = etree.SubElement(body, f"{WNS}sectPr")
    for child_local in spec.input_children:
        child = etree.SubElement(sectPr, f"{WNS}{child_local}")
        _set_child_attrs(child, child_local, matrix.child_attrs)

    dest.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(dest))
    return dest


# CT_SectPr canonical core. Skips relationship-based children
# (headerReference, footerReference, printerSettings) which would need
# corresponding rels in the package, plus sectPrChange (track-change)
# and w15:footnoteColumns (Office 2015 extension).
SECTPR_MATRIX = OrderingMatrix(
    parent_local="sectPr",
    spec_section="ECMA-376 §17.6.18",
    canonical_children=(
        "footnotePr", "endnotePr", "type", "pgSz", "pgMar", "paperSrc",
        "pgBorders", "lnNumType", "pgNumType", "cols", "formProt", "vAlign",
        "noEndnote", "titlePg", "textDirection", "bidi", "rtlGutter", "docGrid",
    ),
    child_attrs=SECTPR_CHILD_ATTRS,
    materialize=materialize_sectpr_scenario,
)


# --- registry --------------------------------------------------------------

MATRICES: dict[str, OrderingMatrix] = {
    "trpr": TRPR_MATRIX,
    "tblpr": TBLPR_MATRIX,
    "tcpr": TCPR_MATRIX,
    "sectpr": SECTPR_MATRIX,
}
